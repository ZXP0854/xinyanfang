"""
第四阶段教程导入（数据收集与分析）
扫描 static/uploads/ 中匹配的文件，按字节大小精确匹配

用法:
  cd /www/wwwroot/xinyanfang/backend && source venv/bin/activate && python import_workflow_4.py
"""
import os, sys
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

import mammoth
from app import create_app
from models import db, Tutorial

# ── 工作流文件映射（文件大小 → node_id → 标题 → 类型） ──
WORKFLOW_FILES = [
    # ═══ 4-1 变量量表选择 ═══
    (165751111, '4-1',   '变量量表选择',                         'mp4'),
    (10496514,  '4-1',   '变量量表选择',                         'docx'),

    # ═══ 4-2-1 问卷类的数据收集（2视频+3文档） ═══
    (74698070,  '4-2-1', '问卷类的数据收集',                      'mp4'),
    (140840422, '4-2-1', '问卷类的数据收集',                      'mp4'),
    (17533,     '4-2-1', '问卷类的数据收集',                      'docx'),
    (5152443,   '4-2-1', '问卷类的数据收集',                      'docx'),
    (9509431,   '4-2-1', '问卷类的数据收集',                      'docx'),

    # ═══ 4-2-2 行为实验数据收集—Psychopy ═══
    (329855183, '4-2-2', '行为实验数据收集—Psychopy',             'mp4'),
    (9720131,   '4-2-2', '行为实验数据收集—Psychopy',             'docx'),

    # ═══ 4-2-3 近红外脑成像数据收集（3视频+3文档） ═══
    (59446459,  '4-2-3', '近红外脑成像数据收集',                  'mp4'),
    (89505475,  '4-2-3', '近红外脑成像数据收集',                  'mp4'),
    (37717310,  '4-2-3', '近红外脑成像数据收集',                  'mp4'),
    (7668880,   '4-2-3', '近红外脑成像数据收集',                  'docx'),
    (747916,    '4-2-3', '近红外脑成像数据收集',                  'docx'),
    (1324290,   '4-2-3', '近红外脑成像数据收集',                  'docx'),

    # ═══ 4-3-1 数据预处理及数据转换 ═══
    (39891923,  '4-3-1', '数据预处理及数据转换',                  'mp4'),
    (1718054,   '4-3-1', '数据预处理及数据转换',                  'docx'),

    # ═══ 4-3-2 数据质量筛查与极端值处理 ═══
    (146861811, '4-3-2', '数据质量筛查与极端值处理',             'mp4'),
    (2016708,   '4-3-2', '数据质量筛查与极端值处理',             'docx'),

    # ═══ 4-4-1 描述性统计与差异检验（2视频+2文档） ═══
    (179800094, '4-4-1', '描述性统计与差异检验',                  'mp4'),
    (44942147,  '4-4-1', '描述性统计与差异检验',                  'mp4'),
    (1539045,   '4-4-1', '描述性统计与差异检验',                  'docx'),
    (360500,    '4-4-1', '描述性统计与差异检验',                  'docx'),

    # ═══ 4-4-2 相关分析与回归分析 ═══
    (63718765,  '4-4-2', '相关分析与回归分析',                    'mp4'),
    (1187434,   '4-4-2', '相关分析与回归分析',                    'docx'),

    # ═══ 4-5-1 SPSS中介与调节效应分析 ═══
    (301952278, '4-5-1', 'SPSS中介与调节效应分析',               'mp4'),
    (1904745,   '4-5-1', 'SPSS中介与调节效应分析',               'docx'),

    # ═══ 4-5-2 高级统计分析方法 ═══
    (178223267, '4-5-2', '高级统计分析方法——Mplus模型和R语言',   'mp4'),
    (344939,    '4-5-2', '高级统计分析方法——Mplus模型和R语言',   'docx'),
]

uploads_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'static', 'uploads')
app = create_app('production')


def find_file(size, ext):
    """按文件大小+扩展名精确匹配"""
    if not os.path.isdir(uploads_dir):
        return None
    for fname in os.listdir(uploads_dir):
        fpath = os.path.join(uploads_dir, fname)
        if not os.path.isfile(fpath):
            continue
        if os.path.getsize(fpath) == size and os.path.splitext(fname)[1].lower() == ('.' + ext if not ext.startswith('.') else ext):
            return fpath
    return None


def convert_docx(filepath):
    """DOCX → HTML，mammoth 失败时用 python-docx 提取文本"""
    try:
        with open(filepath, 'rb') as f:
            result = mammoth.convert_to_html(f)
        html = result.value.replace('src="images/', 'src="/static/uploads/images/')
        for m in (result.messages or []):
            msg = m.message if hasattr(m, 'message') else str(m)
            if msg:
                print(f'    [mammoth] {msg}')
        if html.strip():
            return html
        print(f'    [WARN] mammoth返回空内容，尝试python-docx提取')
    except Exception as e:
        print(f'    [WARN] mammoth失败: {e}，尝试python-docx提取')

    # python-docx 兜底：提取所有段落，保留标题样式
    try:
        import docx
        doc = docx.Document(filepath)
        parts = []
        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue
            style_name = (p.style.name if p.style and p.style.name else '').lower()
            if 'heading 1' in style_name or '标题 1' in style_name:
                parts.append(f'<h4>{text}</h4>')
            elif 'heading 2' in style_name or '标题 2' in style_name:
                parts.append(f'<p><strong>{text}</strong></p>')
            elif 'heading' in style_name or '标题' in style_name:
                parts.append(f'<p><strong>{text}</strong></p>')
            else:
                parts.append(f'<p>{text}</p>')

        # 表格提取
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = ''.join(f'<td>{cell.text}</td>' for cell in row.cells)
                rows.append(f'<tr>{cells}</tr>')
            parts.append(f'<table>{"".join(rows)}</table>')

        html = '\n'.join(parts)
        if html.strip():
            print(f'    [OK] python-docx提取成功 ({len(parts)} 个元素)')
            return html
        else:
            print(f'    [FAIL] python-docx也未提取到内容')
            return ''
    except Exception as e2:
        print(f'    [FAIL] python-docx也失败: {e2}')
        return ''


with app.app_context():
    node_docx = {}
    node_videos = {}
    node_failed_docx = {}
    node_titles = {}

    found_docx = 0
    found_videos = 0
    skipped = 0

    for size, node_id, title, ftype in WORKFLOW_FILES:
        ext_map = {'docx': '.docx', 'mp4': '.mp4', 'mov': '.mov', 'pdf': '.pdf'}
        ext = ext_map.get(ftype, '.' + ftype)
        filepath = find_file(size, ext.lstrip('.'))
        if not filepath:
            print(f'[SKIP] {node_id} {title} ({ftype}, {size} bytes) — 未找到')
            skipped += 1
            continue

        basename = os.path.basename(filepath)
        node_titles[node_id] = title

        if ftype == 'docx':
            print(f'[DOCX] {node_id} {title} ← {basename}')
            html = convert_docx(filepath)
            if html.strip():
                node_docx.setdefault(node_id, []).append(html)
                found_docx += 1
            else:
                node_failed_docx.setdefault(node_id, []).append(basename)
                print(f'    [FALLBACK] 转换失败，将提供下载链接')
        elif ftype in ('mp4', 'mov'):
            print(f'[{ftype.upper()}]  {node_id} {title} ← {basename}')
            node_videos.setdefault(node_id, []).append((basename, ftype == 'mov'))
            found_videos += 1

    print(f'\n找到 {found_docx} 个 DOCX + {found_videos} 个视频，跳过 {skipped}')

    updated = 0
    created = 0

    for node_id, title in sorted(node_titles.items()):
        parts = []

        video_list = node_videos.get(node_id, [])
        docx_list = node_docx.get(node_id, [])

        video_count = len(video_list)
        docx_count = len(docx_list)

        if video_count == docx_count and video_count >= 1:
            for i in range(video_count):
                v, is_mov = video_list[i]
                parts.append(
                    f'<video controls preload="metadata" '
                    f'style="width:100%;max-width:100%;border-radius:12px;display:block;margin-bottom:12px;" '
                    f'poster="/static/uploads/{v.replace(".mp4","").replace(".mov","")}.jpg">'
                    f'<source src="/static/uploads/{v}" type="video/mp4">'
                    f'</video>'
                )
                if docx_count > 1:
                    parts.append(f'<h4>第{i+1}部分</h4>')
                parts.append(docx_list[i])
                if i < video_count - 1:
                    parts.append('<div class="rich-divider"></div>')
        else:
            if video_list:
                vtags = ''.join(
                    f'<video controls preload="metadata" '
                    f'style="width:100%;max-width:100%;border-radius:12px;display:block;margin-bottom:12px;" '
                    f'poster="/static/uploads/{v.replace(".mp4","").replace(".mov","")}.jpg">'
                    f'<source src="/static/uploads/{v}" type="video/mp4">'
                    f'</video>'
                    for v, is_mov in video_list
                )
                parts.append(
                    f'<h4><i class="fa-solid fa-play"></i> 配套视频教程</h4>'
                    f'{vtags}'
                    f'<div class="rich-divider"></div>'
                )
            if docx_list:
                for i, html in enumerate(docx_list):
                    if len(docx_list) > 1:
                        parts.append(f'<h4>第{i+1}部分</h4>')
                    parts.append(html)

        # 转换失败的 DOCX 提供下载链接
        failed_list = node_failed_docx.get(node_id, [])
        if failed_list:
            f_links = ''.join(
                f'<li><a href="/static/uploads/{f}" target="_blank" class="text-link">'
                f'<i class="fa-solid fa-file-word"></i> {f}</a></li>'
                for f in failed_list
            )
            parts.append(
                f'<div class="rich-divider"></div>'
                f'<h4><i class="fa-solid fa-download"></i> 辅助文档（下载查看）</h4>'
                f'<ul>{f_links}</ul>'
            )

        full_html = '\n'.join(parts)
        if not full_html.strip():
            print(f'[WARN] {node_id} {title} — 无有效内容')
            continue

        if '<div class="tutorial-rich"' not in full_html:
            full_html = (
                f'<div class="tutorial-rich"><div class="rich-divider"></div>'
                f'<div class="rich-block"><div class="rich-text">{full_html}</div></div></div>'
            )

        t = Tutorial.query.filter_by(node_id=node_id, category='workflow').first()
        if t:
            t.content = full_html
            t.summary = title
            updated += 1
            print(f'[UPDATE] {node_id} {title} ({len(full_html)} chars)')
        else:
            t = Tutorial(
                node_id=node_id,
                title=title,
                summary=title,
                content=full_html,
                category='workflow',
                is_published=True,
            )
            db.session.add(t)
            created += 1
            print(f'[CREATE] {node_id} {title} ({len(full_html)} chars)')

    db.session.commit()

    print(f'\n{"=" * 55}')
    print(f'  更新: {updated}  新建: {created}  跳过: {skipped}')
    print(f'  工作流教程总数: {Tutorial.query.filter_by(category="workflow").count()}')
    print(f'{"=" * 55}')
